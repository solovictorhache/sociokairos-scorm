#!/usr/bin/env python3
"""Valida que el .docx generado por la exportación Word tiene la forma
esperada: al menos 15 secciones numeradas y al menos 2 tablas
(operacionalización + libro de códigos). Uso: validate_docx.py <ruta.docx>
"""
import sys
from docx import Document


def main(path):
    d = Document(path)
    headings = [
        p.text for p in d.paragraphs
        if p.runs and any(r.bold for r in p.runs)
        and len(p.text) < 100
        and p.text[:2].rstrip(".").isdigit()
    ]
    assert len(headings) >= 15, f"se esperaban al menos 15 secciones numeradas, hay {len(headings)}: {headings}"
    assert len(d.tables) >= 2, f"se esperaban al menos 2 tablas, hay {len(d.tables)}"
    print(f"OK: {len(headings)} secciones, {len(d.tables)} tablas en {path}")


if __name__ == "__main__":
    main(sys.argv[1])
